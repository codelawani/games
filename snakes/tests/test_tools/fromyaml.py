#!/usr/bin/env python3
import yaml
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define the table structure (number of rows and columns)
ROWS = 1
COLUMNS = 7
file = 'snakes_test.yaml'
# Create a new Word document
doc = docx.Document()

# Add a table to the document with the specified number of rows and columns
table = doc.add_table(rows=ROWS, cols=COLUMNS)

# Set table style
table.style = 'Table Grid'

# Set table alignment
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Define the header row
header = table.rows[0].cells
header[0].text = 'ID'
header[1].text = 'Description'
header[2].text = 'Steps'
header[3].text = 'Expected'
header[4].text = 'Actual'
header[5].text = 'Result'
header[6].text = 'Comment'

# Read the test data from a file
with open(file, 'r') as f:
    tests = yaml.safe_load(f)

# Add the test data to the table
for test in tests:
    row_cells = table.add_row().cells
    row_cells[0].text = str(test['Id'])
    row_cells[1].text = test['Description']
    # Combine multiple steps into one cell
    # row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for step in test['Steps']:
        row_cells[2].text += step + '\n'
    row_cells[3].text = str(test['Expected'])
    row_cells[4].text = str(test['Actual'])
    row_cells[5].text = str(test['Result'])
    row_cells[6].text = str(test['Comment'])

# Save the document
doc.save('test_snakes.docx')
